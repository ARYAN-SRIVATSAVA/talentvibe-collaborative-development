from flask import Flask, jsonify, request, send_from_directory
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import text
from datetime import datetime
import json
import hashlib
import io
import os
import openai
import fitz
import docx
import psutil
import gc
import textract
import threading
import time
from collections import deque
from datetime import datetime, timedelta

# Initialize Flask app
app = Flask(__name__, static_folder='frontend_build', static_url_path='')
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('SQLALCHEMY_DATABASE_URI', 'sqlite:///resumes.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = os.environ.get('SQLALCHEMY_TRACK_MODIFICATIONS', 'false').lower() == 'true'

CORS(app)
db = SQLAlchemy(app)

# Configure OpenAI with environment variable
openai.api_key = os.environ.get('OPENAI_API_KEY', "your-openai-api-key-here")

# Database Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    jobs = db.relationship('Job', backref='user', lazy=True)

class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.Text, nullable=False)
    resumes = db.relationship('Resume', backref='job', lazy=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    candidate_name = db.Column(db.String(120), nullable=True)
    content = db.Column(db.Text, nullable=False)
    content_hash = db.Column(db.String(64), nullable=False, index=True)
    analysis = db.Column(db.Text, nullable=True)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)

    __table_args__ = (db.UniqueConstraint('job_id', 'filename', name='_job_filename_uc'),
                      db.UniqueConstraint('job_id', 'content_hash', name='_job_hash_uc'))

# Interview Management Models
class Interview(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Interview Details
    title = db.Column(db.String(200), nullable=False)
    interview_type = db.Column(db.String(50), nullable=False)  # 'phone', 'video', 'onsite', 'technical'
    duration_minutes = db.Column(db.Integer, default=60)
    status = db.Column(db.String(20), default='scheduled')  # 'scheduled', 'completed', 'cancelled', 'rescheduled'
    
    # Scheduling
    scheduled_at = db.Column(db.DateTime, nullable=True)
    location = db.Column(db.String(200), nullable=True)  # For onsite interviews
    video_link = db.Column(db.String(500), nullable=True)  # For video interviews
    
    # Interviewers
    primary_interviewer = db.Column(db.String(100), nullable=True)
    additional_interviewers = db.Column(db.Text, nullable=True)  # JSON array of interviewer names
    
    # Notes
    pre_interview_notes = db.Column(db.Text, nullable=True)
    post_interview_notes = db.Column(db.Text, nullable=True)
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='interviews')
    job = db.relationship('Job', backref='interviews')
    user = db.relationship('User', backref='interviews')

class InterviewFeedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    interview_id = db.Column(db.Integer, db.ForeignKey('interview.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Overall Assessment
    overall_rating = db.Column(db.Integer, nullable=False)  # 1-5 scale
    hire_recommendation = db.Column(db.String(20), nullable=False)  # 'strong_hire', 'hire', 'weak_hire', 'no_hire', 'strong_no_hire'
    
    # Detailed Ratings (JSON)
    technical_skills = db.Column(db.Integer, nullable=True)  # 1-5
    communication_skills = db.Column(db.Integer, nullable=True)  # 1-5
    problem_solving = db.Column(db.Integer, nullable=True)  # 1-5
    cultural_fit = db.Column(db.Integer, nullable=True)  # 1-5
    experience_relevance = db.Column(db.Integer, nullable=True)  # 1-5
    
    # Written Feedback
    strengths = db.Column(db.Text, nullable=True)
    areas_of_concern = db.Column(db.Text, nullable=True)
    additional_notes = db.Column(db.Text, nullable=True)
    
    # Interview Questions & Responses
    questions_asked = db.Column(db.Text, nullable=True)  # JSON array of questions
    candidate_responses = db.Column(db.Text, nullable=True)  # JSON array of responses
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref='interview_feedbacks')

class InterviewQuestion(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Question Details
    question_text = db.Column(db.Text, nullable=False)
    question_type = db.Column(db.String(50), nullable=False)  # 'technical', 'behavioral', 'situational', 'culture_fit'
    difficulty = db.Column(db.String(20), default='medium')  # 'easy', 'medium', 'hard'
    category = db.Column(db.String(100), nullable=True)  # e.g., 'algorithms', 'system_design', 'leadership'
    
    # Usage Tracking
    times_used = db.Column(db.Integer, default=0)
    avg_rating = db.Column(db.Float, default=0.0)  # Average rating from interviewers
    
    # Timestamps
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    job = db.relationship('Job', backref='interview_questions')
    user = db.relationship('User', backref='interview_questions')

# Database connection management
def init_database():
    """Initialize database with connection pooling"""
    try:
        with app.app_context():
            db.create_all()
            print("Database initialized successfully")
    except Exception as e:
        print(f"Database initialization error: {e}")

def get_db_session():
    """Get a database session with error handling"""
    try:
        return db.session
    except Exception as e:
        print(f"Database session error: {e}")
        # Try to recreate session
        db.session.remove()
        return db.session

# Initialize database
init_database()

# API Rate Limiting and Error Handling
import time
import threading
from collections import deque
from datetime import datetime, timedelta

# Rate limiting configuration
RATE_LIMIT_REQUESTS = 20  # Increased from 10 to 20 requests per minute
RATE_LIMIT_WINDOW = 60   # seconds
api_request_times = deque()
api_lock = threading.Lock()

def check_rate_limit():
    """Check if we're within rate limits with better distribution"""
    with api_lock:
        now = datetime.now()
        # Remove old requests outside the window
        while api_request_times and (now - api_request_times[0]) > timedelta(seconds=RATE_LIMIT_WINDOW):
            api_request_times.popleft()
        
        # Check if we can make a request
        if len(api_request_times) >= RATE_LIMIT_REQUESTS:
            return False
        
        # Add current request
        api_request_times.append(now)
        return True

def wait_for_rate_limit():
    """Wait until we can make another API request with exponential backoff"""
    wait_time = 10
    while not check_rate_limit():
        print(f"Rate limit reached, waiting {wait_time} seconds...")
        time.sleep(wait_time)
        wait_time = min(wait_time * 2, 60)  # Exponential backoff, max 60 seconds

def extract_text_from_doc_binary(file_stream):
    """Extract text from .doc binary file without external dependencies"""
    try:
        # Convert binary to string and look for readable text
        # .doc files have text embedded in specific patterns
        binary_str = str(file_stream)
        
        # Look for common text patterns in .doc files
        text_patterns = []
        
        # Try to find text between common delimiters
        import re
        
        # Look for readable text sequences
        # This is a simplified approach that works for most .doc files
        text_chunks = re.findall(r'[A-Za-z0-9\s\.\,\;\:\!\?\-\(\)\[\]\{\}\@\#\$\%\&\*\+\=\_\|\~\`\'\"]{10,}', binary_str)
        
        if text_chunks:
            # Join the text chunks and clean up
            content = ' '.join(text_chunks)
            
            # Remove excessive whitespace and clean up
            content = re.sub(r'\s+', ' ', content)
            content = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\{\}\@\#\$\%\&\*\+\=\_\|\~\`\'\"]', ' ', content)
            content = content.strip()
            
            # Limit length
            if len(content) > 50000:
                content = content[:50000]
                
            return content
        else:
            # If no readable text found, return a placeholder
            return f"Document content could not be extracted. Please review manually."
            
    except Exception as e:
        print(f"Error extracting text from .doc binary: {e}")
        return f"Document: Content extraction failed. Please review manually."

def analyze_resume_with_ai(job_description, resume_text, filename):
    """Analyze resume using OpenAI GPT-4 with comprehensive error handling"""
    
    print(f"Starting AI analysis for {filename}")
    print(f"OpenAI API key configured: {bool(openai.api_key and openai.api_key != 'your-openai-api-key-here')}")
    print(f"Rate limit check: {check_rate_limit()}")
    
    # Check if OpenAI API key is configured
    if not openai.api_key or openai.api_key == "your-openai-api-key-here":
        print(f"OpenAI API key not configured for {filename}, using fallback analysis")
        return create_fallback_analysis(filename, "OpenAI API key not configured")
    
    # Check rate limit before making API call
    if not check_rate_limit():
        print(f"Rate limit reached for {filename}, using fallback analysis")
        return create_fallback_analysis(filename, "Rate limit reached")
    
    try:
        # Prepare prompt with better error handling
        prompt = f"""
You are an expert HR recruiter and technical interviewer. Analyze this resume against the job description and provide a comprehensive assessment.

Job Description:
{job_description[:1000]}

Resume Content:
{resume_text[:2000]}

Please provide your analysis in the following JSON format:
{{
    "candidate_name": "Extract name from resume or 'Name Not Found'",
    "fit_score": 85,
    "bucket": "⚡ Book-the-Call",
    "reasoning": "Brief explanation of why this candidate fits or doesn't fit",
    "summary_points": [
        "Key strength 1",
        "Key strength 2",
        "Key concern 1"
    ],
    "skill_matrix": {{
        "matches": ["Skill 1", "Skill 2"],
        "gaps": ["Missing skill 1", "Missing skill 2"]
    }},
    "timeline": [
        {{
            "period": "2022-2024",
            "role": "Software Engineer",
            "company": "Tech Company",
            "details": "Brief description of role and achievements"
        }}
    ],
    "logistics": {{
        "compensation": "Expected salary range or 'Not specified'",
        "notice_period": "Notice period or 'Not specified'",
        "work_authorization": "Visa status or 'Not specified'",
        "location": "Preferred location or 'Not specified'"
    }}
}}

Focus on:
1. Technical skills alignment
2. Experience relevance
3. Cultural fit indicators
4. Red flags or concerns
5. Overall hire recommendation

Return only valid JSON.
"""
        
        print(f"Making OpenAI API call for {filename}")
        
        # Make API call with timeout and retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                print(f"OpenAI API attempt {attempt + 1}/{max_retries} for {filename}")
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                    max_tokens=1000
                )
                
                analysis_text = response.choices[0].message.content.strip()
                print(f"OpenAI API response received for {filename}: {len(analysis_text)} characters")
                
                # Validate JSON response
                try:
                    parsed_json = json.loads(analysis_text)
                    print(f"AI analysis successful for {filename} - Score: {parsed_json.get('fit_score', 'N/A')}")
                    return analysis_text
                except json.JSONDecodeError as e:
                    print(f"Invalid JSON response for {filename}: {e}")
                    print(f"Response content: {analysis_text[:200]}...")
                    return create_fallback_analysis(filename, "Invalid JSON response")
                    
            except openai.RateLimitError as e:
                print(f"Rate limit error for {filename}, attempt {attempt + 1}/{max_retries}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(60)  # Wait 1 minute before retry
                    continue
                else:
                    return create_fallback_analysis(filename, "Rate limit exceeded")
                    
            except openai.QuotaExceededError as e:
                print(f"Quota exceeded for {filename}: {e}")
                return create_fallback_analysis(filename, "API quota exceeded")
                
            except openai.AuthenticationError as e:
                print(f"Authentication error for {filename}: {e}")
                return create_fallback_analysis(filename, "API authentication failed")
                
            except openai.APIConnectionError as e:
                print(f"API connection error for {filename}, attempt {attempt + 1}/{max_retries}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(30)  # Wait 30 seconds before retry
                    continue
                else:
                    return create_fallback_analysis(filename, "API connection failed")
                    
            except Exception as e:
                print(f"Unexpected error for {filename}, attempt {attempt + 1}/{max_retries}: {e}")
                if attempt < max_retries - 1:
                    time.sleep(30)
                    continue
                else:
                    return create_fallback_analysis(filename, f"Unexpected error: {str(e)}")
    
    except Exception as e:
        print(f"Critical error in AI analysis for {filename}: {e}")
        return create_fallback_analysis(filename, f"Critical error: {str(e)}")

def create_fallback_analysis(filename, error_reason):
    """Create a fallback analysis when AI fails"""
    return json.dumps({
        "candidate_name": filename.split('.')[0].replace('_', ' '),
        "fit_score": 70,
        "bucket": "🛠️ Bench Prospect",
        "reasoning": f"AI analysis unavailable: {error_reason}. Manual review recommended for detailed evaluation.",
        "summary_points": [
            "Analysis completed with fallback",
            "AI service temporarily unavailable",
            "Manual review recommended"
        ],
        "skill_matrix": {"matches": [], "gaps": []},
        "timeline": [],
        "logistics": {
            "compensation": "Not specified",
            "notice_period": "Not specified", 
            "work_authorization": "Not specified",
            "location": "Not specified"
        }
    })
@app.route("/", methods=["GET", "POST"])
def root():
    """Handle root requests - serve React app for GET, redirect POST to analyze"""
    if request.method == "GET":
        # Serve the React app
        try:
            if not os.path.exists(app.static_folder):
                return f"""<h1>TalentVibe</h1><p>Static folder not found: {app.static_folder}</p>"""
            
            index_path = os.path.join(app.static_folder, "index.html")
            if not os.path.exists(index_path):
                return f"""<h1>TalentVibe</h1><p>index.html not found in: {app.static_folder}</p>"""
            
            return send_from_directory(app.static_folder, "index.html")
        except Exception as e:
            return f"""<h1>TalentVibe</h1><p>Error serving React app: {str(e)}</p><p>Static folder: {app.static_folder}</p>"""
    elif request.method == "POST":
        # Redirect POST requests to the analyze endpoint
        return jsonify({"error": "Please use /api/analyze endpoint for POST requests"}), 405
@app.route("/<path:path>")
def serve_static(path):
    """Serve static files from React build"""
    try:
        # Check if the file exists
        file_path = os.path.join(app.static_folder, path)
        if not os.path.exists(file_path):
            return f"File not found: {path}", 404
        
        return send_from_directory(app.static_folder, path)
    except Exception as e:
        return f"Error serving {path}: {str(e)}", 500

@app.route('/api/health')
def health_check():
    """Simple health check without external dependencies"""
    try:
        # Test database connection only
        db.session.execute(text('SELECT 1'))
        return jsonify({
            'status': 'healthy',
            'message': 'TalentVibe API is running',
            'database': 'connected'
        })
    except Exception as e:
        return jsonify({
            'status': 'degraded',
            'message': 'Database connection issue',
            'database': f'error: {str(e)}'
        }), 500

@app.route('/api/test')
def test_endpoint():
    return jsonify({'message': 'API test successful'})

# Resource management and monitoring

def check_system_resources():
    """Check if system has enough resources"""
    try:
        memory = psutil.virtual_memory()
        cpu_percent = psutil.cpu_percent(interval=1)
        
        # Check memory usage
        if memory.percent > 90:
            return False, f"Memory usage too high: {memory.percent}%"
        
        # Check CPU usage
        if cpu_percent > 90:
            return False, f"CPU usage too high: {cpu_percent}%"
        
        return True, f"Memory: {memory.percent}%, CPU: {cpu_percent}%"
    except ImportError as e:
        print(f"psutil not available: {e}")
        return True, "Resource monitoring unavailable"
    except Exception as e:
        print(f"Resource check error: {e}")
        return True, "Resource check unavailable"

def cleanup_resources():
    """Clean up resources to prevent memory leaks"""
    try:
        # Force garbage collection
        gc.collect()
        
        # Don't clear the session here - it causes binding issues
        # db.session.remove()  # REMOVED - causes session binding errors
        
        print("Resource cleanup completed")
    except Exception as e:
        print(f"Resource cleanup error: {e}")

def refresh_session():
    """Refresh the database session safely"""
    try:
        db.session.commit()
        db.session.flush()
    except Exception as e:
        print(f"Session refresh error: {e}")
        # Only remove session if there's a critical error
        db.session.rollback()


def extract_text_from_file(file_stream):
    """Extract text from various file formats"""
    try:
        filename = file_stream.filename.lower()
        file_stream.seek(0)
        
        if filename.endswith(".pdf"):
            try:
                import fitz
                doc = fitz.open(stream=file_stream.read(), filetype="pdf")
                content = ""
                for page in doc:
                    content += page.get_text()
                doc.close()
                return content[:50000]
            except Exception as e:
                print(f"PDF extraction error: {e}")
                return f"PDF extraction failed for {filename}"
        elif filename.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(file_stream)
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                return content[:50000]
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                return f"DOCX extraction failed for {filename}"
        elif filename.endswith(".doc"):
            return extract_text_from_doc_binary(file_stream)
        elif filename.endswith(".txt"):
            return file_stream.read().decode("utf-8")[:50000]
        else:
            return f"Unsupported file type: {filename}"
    except Exception as e:
        print(f"File extraction error: {e}")
        return f"File extraction failed: {str(e)}"

@app.route("/api/jd/check", methods=["POST"])
def check_existing_jd():
    """Check if a job description file already exists in the database"""
    try:
        if "jd_file" not in request.files:
            return jsonify({"error": "No job description file provided"}), 400
        
        jd_file = request.files["jd_file"]
        if jd_file.filename == "":
            return jsonify({"error": "No file selected"}), 400
        
        # Extract content from the JD file
        try:
            content = extract_text_from_file(jd_file)
            content_hash = hashlib.sha256(content.encode()).hexdigest()
        except Exception as e:
            print(f"Error extracting JD content: {e}")
            return jsonify({"error": "Could not read job description file"}), 400
        
        # Check if this JD content already exists
        existing_job = Job.query.filter_by(description=content).first()
        
        if existing_job:
            # Get resume count for this job
            resume_count = Resume.query.filter_by(job_id=existing_job.id).count()
            
            return jsonify({
                "exists": True,
                "message": f"Job description already exists (Job #{existing_job.id})",
                "job": {
                    "id": existing_job.id,
                    "resume_count": resume_count
                },
                "jd_file": {
                    "filename": jd_file.filename,
                    "content": content,
                    "file_type": jd_file.filename.split(".")[-1] if "." in jd_file.filename else "unknown",
                    "created_at": existing_job.id  # Using job ID as proxy for creation time
                }
            })
        else:
            return jsonify({
                "exists": False,
                "message": "This is a new job description"
            })
            
    except Exception as e:
        print(f"JD check error: {e}")
        return jsonify({"error": "Internal server error"}), 500

@app.route("/api/resumes/check-duplicates", methods=["POST"])
def check_resume_duplicates():
    """Check for duplicate resumes within a specific job"""
    try:
        if "job_id" not in request.form:
            return jsonify({"error": "Job ID is required"}), 400
        
        job_id = int(request.form["job_id"])
        resume_files = request.files.getlist("resumes")
        
        if not resume_files:
            return jsonify({"error": "No resume files provided"}), 400
        
        # Get existing resumes for this job
        existing_resumes = Resume.query.filter_by(job_id=job_id).all()
        existing_hashes = {resume.content_hash for resume in existing_resumes}
        
        duplicates = []
        unique_count = 0
        total_files = len(resume_files)
        
        for resume_file in resume_files:
            if resume_file.filename == "":
                continue
                
            try:
                content = extract_text_from_file(resume_file)
                content_hash = hashlib.sha256(content.encode()).hexdigest()
                
                # Check if this content already exists
                existing_resume = Resume.query.filter_by(job_id=job_id, content_hash=content_hash).first()
                
                if existing_resume:
                    duplicates.append({
                        "filename": resume_file.filename,
                        "duplicate_of": {
                            "resume_filename": existing_resume.filename,
                            "candidate_name": existing_resume.candidate_name,
                            "job_id": existing_resume.job_id
                        }
                    })
                else:
                    unique_count += 1
                    
            except Exception as e:
                duplicates.append({
                    "filename": resume_file.filename,
                    "error": f"Could not process file: {str(e)}"
                })
        
        return jsonify({
            "job_id": job_id,
            "total_files": total_files,
            "unique_count": unique_count,
            "duplicate_count": len(duplicates),
            "duplicates": duplicates
        })
        
    except Exception as e:
        print(f"Duplicate check error: {e}")
        return jsonify({"error": "Internal server error"}), 500
@app.route("/api/analyze", methods=["POST"])
def analyze_resumes():
    """Analyze uploaded resumes with AI - return immediately, process asynchronously"""
    try:
        print("=== Starting analyze_resumes request ===")
        
        # Check request size
        content_length = request.content_length
        if content_length and content_length > 100 * 1024 * 1024:  # 100MB limit
            print(f"Request too large: {content_length} bytes")
            return jsonify({'error': 'Request too large. Maximum 100MB allowed.'}), 413
        
        # Check system resources before starting
        resources_ok, resource_status = check_system_resources()
        if not resources_ok:
            print(f"Resource warning: {resource_status}")
            # Continue but with caution
        
        # Get or create default user with proper session management
        try:
            user = User.query.filter_by(username='default_user').first()
            if not user:
                user = User(username='default_user')
                db.session.add(user)
                db.session.commit()
                print("Created default user")
            else:
                print("Found existing default user")
        except Exception as e:
            print(f"User creation error: {e}")
            db.session.rollback()
            return jsonify({'error': 'Database session error'}), 500
        
        if 'job_description' not in request.form:
            print("No job description provided")
            return jsonify({'error': 'Job description is required'}), 400
        
        job_description = request.form['job_description']
        print(f"Job description length: {len(job_description)}")
        
        # Process job description files if they exist
        if 'job_description_files' in request.files:
            job_description_files = request.files.getlist('job_description_files')
            if job_description_files and any(f.filename for f in job_description_files):
                print(f"Processing {len(job_description_files)} job description files")
                jd_content = []
                
                for jd_file in job_description_files:
                    if jd_file.filename:
                        try:
                            # Extract text from job description file
                            content = extract_text_from_file(jd_file)
                            jd_content.append(f"File: {jd_file.filename}\n{content}")
                            print(f"Extracted {len(content)} characters from {jd_file.filename}")
                        except Exception as e:
                            print(f"Error extracting text from {jd_file.filename}: {e}")
                            jd_content.append(f"File: {jd_file.filename}\nError extracting content: {str(e)}")
                
                if jd_content:
                    # Use the extracted content instead of the placeholder text
                    job_description = "\n\n".join(jd_content)
                    print(f"Updated job description length: {len(job_description)}")
        
        # Create job with proper session management
        try:
            job = Job(description=job_description, user_id=user.id)
            db.session.add(job)
            db.session.commit()
            job_id = job.id
            print(f"Created job with ID: {job_id}")
        except Exception as e:
            print(f"Job creation error: {e}")
            db.session.rollback()
            return jsonify({'error': 'Failed to create job'}), 500
        
        # Get uploaded files
        resume_files = request.files.getlist('resumes')
        print(f"Received {len(resume_files)} resume files")
        
        if not resume_files or all(f.filename == '' for f in resume_files):
            print("No valid resume files uploaded")
            return jsonify({'error': 'No resume files uploaded'}), 400
        
        # Limit number of files to prevent resource exhaustion
        max_files = 30  # Increased from 10 to 30
        if len(resume_files) > max_files:
            print(f"Too many files: {len(resume_files)} > {max_files}")
            return jsonify({
                'error': f'Too many files. Maximum {max_files} files allowed per request.'
            }), 400
        
        # Store files temporarily and return immediately
        file_data = []
        total_size = 0
        for resume_file in resume_files:
            if resume_file.filename:
                try:
                    file_content = resume_file.read()
                    file_size = len(file_content)
                    total_size += file_size
                    file_data.append({
                        'filename': resume_file.filename,
                        'content': file_content,
                        'job_id': job_id
                    })
                    print(f"Stored file: {resume_file.filename} ({file_size} bytes)")
                except Exception as e:
                    print(f"Error reading file {resume_file.filename}: {e}")
        
        print(f"Successfully stored {len(file_data)} files for background processing (total size: {total_size} bytes)")
        
        # Return immediately with job_id
        response_data = {
            'message': f'Analysis queued successfully for {len(file_data)} resume(s). Processing in background.',
            'job_id': job_id,
            'total_files': len(file_data),
            'status': 'queued',
            'redirect_url': f'/jobs/{job_id}'
        }
        
        # Start background processing
        try:
            import threading
            processing_thread = threading.Thread(
                target=process_resumes_background,
                args=(file_data, job_description, job_id)
            )
            processing_thread.daemon = False  # Changed from True to False - prevents thread from being killed
            processing_thread.start()
            print(f"Started background processing thread for job {job_id}")
        except Exception as e:
            print(f"Error starting background thread: {e}")
            # Continue anyway - the job is created and user can check later
        
        print("=== analyze_resumes request completed successfully ===")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Critical error in analyze_resumes: {e}")
        import traceback
        traceback.print_exc()
        db.session.rollback()
        return jsonify({'error': f'Critical error: {str(e)}'}), 500

def process_resumes_background(file_data, job_description, job_id):
    """Process resumes in background thread"""
    processed_files = []
    skipped_files = []
    
    try:
        print(f"Starting background processing for job {job_id} with {len(file_data)} files")
        
        for i, file_info in enumerate(file_data):
            try:
                print(f"=== Processing file {i+1}/{len(file_data)}: {file_info['filename']} ===")
                
                # Check resources every 5 files (files 5, 10, 15, etc.)
                if (i + 1) % 5 == 0:
                    print(f"Resource check at file {i+1}")
                    gc.collect()
                    try:
                        resources_ok, resource_status = check_system_resources()
                        print(f"Resource status: {resource_status}")
                        if not resources_ok:
                            print("Resource limit reached, stopping processing")
                            break
                    except Exception as resource_error:
                        print(f"Resource check failed: {resource_error}")
                        # Continue processing even if resource check fails
                
                filename = file_info['filename']
                file_stream = file_info['content']
                
                print(f"Processing file {i+1}/{len(file_data)}: {filename}")
                
                # Extract text from file with size limits
                content = ""
                try:
                    if filename.lower().endswith('.pdf'):
                        pdf_doc = fitz.open(stream=file_stream, filetype='pdf')
                        for page in pdf_doc:
                            content += page.get_text()
                            if len(content) > 50000:
                                break
                        pdf_doc.close()
                    elif filename.lower().endswith('.docx'):
                        doc = docx.Document(io.BytesIO(file_stream))
                        for para in doc.paragraphs:
                            content += para.text + '\n'
                            if len(content) > 50000:
                                break
                    elif filename.lower().endswith('.doc'):
                        # Handle .doc files with robust fallback
                        try:
                            # Try to use textract if available
                            try:
                                import tempfile
                                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                                    temp_file.write(file_stream)
                                    temp_file_path = temp_file.name
                                
                                content = textract.process(temp_file_path).decode('utf-8')
                                os.unlink(temp_file_path)
                                
                                if len(content) > 50000:
                                    content = content[:50000]
                                    
                            except ImportError:
                                # textract not available, use alternative method
                                print(f"Textract not available for {filename}, using alternative method")
                                content = extract_text_from_doc_binary(file_stream)
                                
                        except Exception as e:
                            print(f"Textract error for {filename}: {e}")
                            # Fallback: try to extract text from binary
                            try:
                                content = extract_text_from_doc_binary(file_stream)
                            except Exception as fallback_error:
                                print(f"Fallback extraction failed for {filename}: {fallback_error}")
                                # Last resort: use filename as content
                                content = f"Document: {filename}\n\nContent extraction failed. Please review manually."
                    else:
                        content = file_stream.decode('utf-8')[:50000]
                except Exception as e:
                    print(f"File extraction error for {filename}: {e}")
                    skipped_files.append({'filename': filename, 'reason': f'File extraction failed: {str(e)}'})
                    continue
                
                if not content.strip():
                    skipped_files.append({'filename': filename, 'reason': 'Empty file'})
                    continue
                
                # Check for duplicates
                try:
                    content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
                    with app.app_context():
                        existing = Resume.query.filter_by(job_id=job_id, content_hash=content_hash).first()
                        if existing:
                            skipped_files.append({'filename': filename, 'reason': 'Duplicate'})
                            continue
                except Exception as e:
                    print(f"Duplicate check error for {filename}: {e}")
                    continue
                
                # Analyze with AI
                try:
                    print(f"Starting AI analysis for {filename}")
                    analysis_text = analyze_resume_with_ai(job_description, content, filename)
                    print(f"AI analysis completed for {filename}")
                    
                    try:
                        analysis_json = json.loads(analysis_text)
                        print(f"JSON parsing successful for {filename}")
                    except json.JSONDecodeError as json_error:
                        print(f"JSON parsing failed for {filename}: {json_error}")
                        analysis_json = {"candidate_name": "Parse Error", "fit_score": 0, "bucket": "Error"}
                    
                    candidate_name = analysis_json.get('candidate_name', 'Not Provided')
                    if candidate_name == 'Name Not Found':
                        candidate_name = filename.split('.')[0].replace('_', ' ')
                    
                    print(f"Processing candidate: {candidate_name}")
                    
                    # Save to database
                    try:
                        print(f"Saving to database: {filename}")
                        with app.app_context():
                            new_resume = Resume(
                                filename=filename,
                                candidate_name=candidate_name,
                                content=content,
                                content_hash=content_hash,
                                analysis=analysis_text,
                                job_id=job_id
                            )
                            db.session.add(new_resume)
                            db.session.commit()
                            processed_files.append(filename)
                            print(f"Successfully processed and saved: {filename}")
                            
                    except Exception as db_error:
                        print(f"Database save error for {filename}: {db_error}")
                        db.session.rollback()
                        # Try fallback save
                        try:
                            print(f"Attempting fallback save for {filename}")
                            fallback_analysis = create_fallback_analysis(filename, "Database save failed")
                            with app.app_context():
                                new_resume = Resume(
                                    filename=filename,
                                    candidate_name=filename.split('.')[0].replace('_', ' '),
                                    content=content,
                                    content_hash=content_hash,
                                    analysis=fallback_analysis,
                                    job_id=job_id
                                )
                                db.session.add(new_resume)
                                db.session.commit()
                                processed_files.append(filename)
                                print(f"Fallback save successful for {filename}")
                        except Exception as fallback_error:
                            print(f"Fallback save failed for {filename}: {fallback_error}")
                            skipped_files.append({'filename': filename, 'reason': 'Database error'})
                    
                except Exception as ai_error:
                    print(f"AI Analysis failed for {filename}: {ai_error}")
                    # Save with fallback analysis
                    try:
                        print(f"Creating fallback analysis for {filename}")
                        fallback_analysis = create_fallback_analysis(filename, str(ai_error))
                        with app.app_context():
                            new_resume = Resume(
                                filename=filename,
                                candidate_name=filename.split('.')[0].replace('_', ' '),
                                content=content,
                                content_hash=content_hash,
                                analysis=fallback_analysis,
                                job_id=job_id
                            )
                            db.session.add(new_resume)
                            db.session.commit()
                            processed_files.append(filename)
                            print(f"Fallback analysis saved for {filename}")
                    except Exception as fallback_error:
                        print(f"Fallback save failed for {filename}: {fallback_error}")
                        skipped_files.append({'filename': filename, 'reason': 'AI and database error'})
                
                # Clean up memory more aggressively
                if "content" in locals(): del content
                if "file_stream" in locals(): del file_stream
                if "analysis_text" in locals(): del analysis_text
                gc.collect()  # Force garbage collection after each file
                
            except Exception as e:
                print(f"File processing error for {filename}: {e}")
                skipped_files.append({'filename': filename, 'reason': str(e)})
                # Clean up memory even on error
                try:
                    if "content" in locals(): del content
                    if "file_stream" in locals(): del file_stream
                    if "analysis_text" in locals(): del analysis_text
                    gc.collect()
                except:
                    pass
        
        # Final cleanup
        gc.collect()
        print(f"Background processing completed for job {job_id}. Processed: {len(processed_files)}, Skipped: {len(skipped_files)}")
        
    except Exception as e:
        print(f"Critical error in background processing for job {job_id}: {e}")
        import traceback
        traceback.print_exc()
        
        # Try to save any processed files even if there was an error
        try:
            print(f"Attempting to save {len(processed_files)} processed files despite error")
            for filename in processed_files:
                print(f"Successfully processed: {filename}")
        except Exception as save_error:
            print(f"Error saving processed files: {save_error}")

@app.route('/api/jobs', methods=['GET'])
def get_jobs():
    """Get all jobs for the default user"""
    user = User.query.filter_by(username='default_user').first()
    if not user:
        return jsonify([])
    
    jobs = Job.query.filter_by(user_id=user.id).order_by(Job.id.desc()).all()
    
    return jsonify([{
        'id': job.id,
        'description': job.description,
        'resume_count': len(job.resumes),
        'created_at': job.id  # Using ID as simple timestamp
    } for job in jobs])

@app.route('/api/jobs/<int:job_id>', methods=['GET'])
def get_job_details(job_id):
    """Get detailed information about a specific job"""
    job = Job.query.get_or_404(job_id)
    
    # Group resumes by analysis result
    resumes_data = []
    for resume in job.resumes:
        try:
            analysis = json.loads(resume.analysis) if resume.analysis else None
        except:
            analysis = None
        
        resumes_data.append({
            'id': resume.id,
            'filename': resume.filename,
            'candidate_name': resume.candidate_name,
            'analysis': analysis,
            'bucket': analysis.get('bucket', 'Unknown') if analysis else 'Unknown',
            'feedback_count': 0  # Mock value
        })
    
    return jsonify({
        'id': job.id,
        'description': job.description,
        'resumes': resumes_data,
        'total_resumes': len(job.resumes)
    })

@app.route('/api/jobs/<int:job_id>', methods=['DELETE'])
def delete_job(job_id):
    """Delete a job and all associated resumes"""
    job = Job.query.get_or_404(job_id)
    
    # Delete associated resumes
    Resume.query.filter_by(job_id=job_id).delete()
    
    # Delete the job
    db.session.delete(job)
    db.session.commit()
    
    return jsonify({'message': 'Job deleted successfully'})

@app.route('/api/data')
def get_data():
    return jsonify({'message': 'TalentVibe API is running'})

# Interview Management Routes
@app.route('/api/interviews', methods=['GET'])
def get_interviews():
    """Get all interviews for the default user"""
    user = User.query.filter_by(username='default_user').first()
    if not user:
        return jsonify([])
    
    interviews = Interview.query.filter_by(user_id=user.id).order_by(Interview.created_at.desc()).all()
    
    interview_data = []
    for interview in interviews:
        interview_data.append({
            'id': interview.id,
            'title': interview.title,
            'interview_type': interview.interview_type,
            'status': interview.status,
            'scheduled_at': interview.scheduled_at.isoformat() if interview.scheduled_at else None,
            'resume_id': interview.resume_id,
            'job_id': interview.job_id,
            'primary_interviewer': interview.primary_interviewer,
            'duration_minutes': interview.duration_minutes,
            'created_at': interview.created_at.isoformat() if interview.created_at else None
        })
    
    return jsonify(interview_data)

@app.route('/api/interviews/<int:resume_id>', methods=['GET'])
def get_resume_interviews(resume_id):
    """Get all interviews for a specific resume"""
    resume = Resume.query.get_or_404(resume_id)
    
    interviews = Interview.query.filter_by(resume_id=resume_id).order_by(Interview.created_at.desc()).all()
    
    interview_data = []
    for interview in interviews:
        interview_data.append({
            'id': interview.id,
            'title': interview.title,
            'interview_type': interview.interview_type,
            'status': interview.status,
            'scheduled_at': interview.scheduled_at.isoformat() if interview.scheduled_at else None,
            'location': interview.location,
            'video_link': interview.video_link,
            'primary_interviewer': interview.primary_interviewer,
            'additional_interviewers': interview.additional_interviewers,
            'duration_minutes': interview.duration_minutes,
            'pre_interview_notes': interview.pre_interview_notes,
            'post_interview_notes': interview.post_interview_notes,
            'created_at': interview.created_at.isoformat() if interview.created_at else None,
            'updated_at': interview.updated_at.isoformat() if interview.updated_at else None
        })
    
    return jsonify({
        'resume_id': resume_id,
        'filename': resume.filename,
        'interviews': interview_data
    })

@app.route('/api/interviews', methods=['POST'])
def create_interview():
    """Create a new interview"""
    data = request.get_json()
    
    if not all(key in data for key in ['resume_id', 'job_id', 'title', 'interview_type']):
        return jsonify({'error': 'Missing required fields'}), 400
    
    # Get or create default user
    user = User.query.filter_by(username='default_user').first()
    if not user:
        user = User(username='default_user')
        db.session.add(user)
        db.session.commit()
    
    interview = Interview(
        resume_id=data['resume_id'],
        job_id=data['job_id'],
        user_id=user.id,
        title=data['title'],
        interview_type=data['interview_type'],
        duration_minutes=data.get('duration_minutes', 60),
        status=data.get('status', 'scheduled'),
        scheduled_at=datetime.fromisoformat(data['scheduled_at']) if data.get('scheduled_at') else None,
        location=data.get('location'),
        video_link=data.get('video_link'),
        primary_interviewer=data.get('primary_interviewer'),
        additional_interviewers=data.get('additional_interviewers'),
        pre_interview_notes=data.get('pre_interview_notes'),
        post_interview_notes=data.get('post_interview_notes')
    )
    
    db.session.add(interview)
    db.session.commit()
    
    return jsonify({
        'message': 'Interview created successfully',
        'interview_id': interview.id
    })

@app.route('/api/interviews/<int:interview_id>', methods=['PUT'])
def update_interview(interview_id):
    """Update an existing interview"""
    interview = Interview.query.get_or_404(interview_id)
    data = request.get_json()
    
    # Update fields if provided
    if 'title' in data:
        interview.title = data['title']
    if 'interview_type' in data:
        interview.interview_type = data['interview_type']
    if 'status' in data:
        interview.status = data['status']
    if 'duration_minutes' in data:
        interview.duration_minutes = data['duration_minutes']
    if 'scheduled_at' in data:
        interview.scheduled_at = datetime.fromisoformat(data['scheduled_at']) if data['scheduled_at'] else None
    if 'location' in data:
        interview.location = data['location']
    if 'video_link' in data:
        interview.video_link = data['video_link']
    if 'primary_interviewer' in data:
        interview.primary_interviewer = data['primary_interviewer']
    if 'additional_interviewers' in data:
        interview.additional_interviewers = data['additional_interviewers']
    if 'pre_interview_notes' in data:
        interview.pre_interview_notes = data['pre_interview_notes']
    if 'post_interview_notes' in data:
        interview.post_interview_notes = data['post_interview_notes']
    
    interview.updated_at = datetime.utcnow()
    db.session.commit()
    
    return jsonify({
        'message': 'Interview updated successfully',
        'interview_id': interview.id
    })

@app.route('/api/interviews/<int:interview_id>', methods=['DELETE'])
def delete_interview(interview_id):
    """Delete an interview"""
    interview = Interview.query.get_or_404(interview_id)
    db.session.delete(interview)
    db.session.commit()
    
    return jsonify({'message': 'Interview deleted successfully'})

# Enhanced error handling and monitoring
import logging
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Error tracking
error_count = 0
last_error_time = None

def log_error(error_type, error_message, context=None):
    """Log errors with context"""
    global error_count, last_error_time
    
    error_count += 1
    last_error_time = datetime.now()
    
    logger.error(f"{error_type}: {error_message}")
    if context:
        logger.error(f"Context: {context}")
    
    # If too many errors, log warning
    if error_count > 10:
        logger.warning(f"High error count: {error_count} errors")

# Global error handler to ensure JSON responses
@app.errorhandler(Exception)
def handle_exception(e):
    log_error("Unhandled Exception", str(e), {
        'endpoint': request.endpoint,
        'method': request.method,
        'url': request.url
    })
    
    import traceback
    traceback.print_exc()
    
    # Return appropriate error response
    if "database" in str(e).lower():
        return jsonify({'error': 'Database connection issue'}), 500
    elif "memory" in str(e).lower():
        return jsonify({'error': 'System resource issue'}), 500
    elif "timeout" in str(e).lower():
        return jsonify({'error': 'Request timeout'}), 408
    else:
        return jsonify({'error': 'Internal server error'}), 500

@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Endpoint not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    log_error("500 Error", str(error))
    return jsonify({'error': 'Internal server error'}), 500

@app.errorhandler(413)
def too_large(error):
    print(f"Request too large error: {error}")
    return jsonify({
        'error': 'Request too large. Please reduce the number of files or file sizes.',
        'details': 'Maximum 100MB total upload size allowed.'
    }), 413

# WSGI application
application = app

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)                                                                 